#!/usr/bin/env python3
"""
Local server for Agentic AI Market Navigator.
Run:
  pip install -r requirements.txt
  python server.py
Then open http://localhost:8000
"""
from __future__ import annotations

import json
import os
import re
import shutil
import sys
import time
from datetime import datetime, timezone
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import urlparse

BASE = Path(__file__).resolve().parent
DATA_DIR = BASE / "data"
UPLOAD_DIR = BASE / "uploads"
USER_UPLOADS = DATA_DIR / "user_uploads.json"
UPLOAD_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)

SECTION_KEYWORDS = {
    "ke-toan": ["kế toán", "tài chính", "hóa đơn", "thuế", "đối soát", "sổ cái", "misa", "cashflow", "invoice"],
    "san-xuat": ["sản xuất", "nhà máy", "bảo trì", "downtime", "qc", "qa", "mes", "iot", "dây chuyền", "phế phẩm"],
    "logistics": ["logistics", "chuỗi cung ứng", "vận tải", "kho", "hải quan", "vận đơn", "packing list", "bill of lading", "tms", "wms"],
    "phap-ly": ["pháp lý", "hợp đồng", "luật", "tuân thủ", "compliance", "legal", "redline", "văn bản pháp luật"],
    "foundation": ["agentic", "ai agent", "llm", "memory", "planner", "tool", "copilot", "rpa", "multi-agent", "reflection", "react"],
    "global": ["global", "world", "mckinsey", "wef", "gartner", "market", "cagr", "billion", "trillion", "sequoia", "a16z"],
    "vietnam": ["việt nam", "vietnam", "fpt", "viettel", "nvidia", "nic", "sme", "doanh nghiệp", "nghị định", "quyết định"],
    "niches": ["ngách", "startup", "mvp", "doanh thu", "khách hàng", "pricing", "saas", "vertical", "mô hình doanh thu"],
    "careers": ["nghề", "career", "job", "workflow architect", "agentops", "evaluator", "governance specialist", "reskilling"],
    "governance": ["risk", "rủi ro", "governance", "iso", "nist", "eu ai act", "bảo mật", "privacy", "prompt injection", "hallucination", "audit"]
}
SUMMARY_KEYWORDS = ["agent", "agentic", "AI", "thị trường", "Việt Nam", "ngách", "nghề", "governance", "rủi ro", "workflow", "SME", "startup", "ROI"]




def parse_content_type(content_type: str) -> tuple[str, dict[str, str]]:
    parts = [p.strip() for p in content_type.split(';')]
    main = parts[0].lower() if parts else ''
    params: dict[str, str] = {}
    for part in parts[1:]:
        if '=' in part:
            k, v = part.split('=', 1)
            params[k.strip().lower()] = v.strip().strip('"')
    return main, params


def parse_content_disposition(value: str) -> dict[str, str]:
    params: dict[str, str] = {}
    for part in value.split(';'):
        part = part.strip()
        if '=' in part:
            k, v = part.split('=', 1)
            params[k.strip().lower()] = v.strip().strip('"')
        else:
            params['type'] = part.lower()
    return params


def read_multipart_file(handler) -> tuple[str, bytes]:
    content_type = handler.headers.get('content-type', '')
    main, params = parse_content_type(content_type)
    if main != 'multipart/form-data' or 'boundary' not in params:
        raise ValueError('Content-Type phải là multipart/form-data')
    length = int(handler.headers.get('content-length', '0'))
    if length <= 0:
        raise ValueError('Request không có nội dung file')
    body = handler.rfile.read(length)
    boundary = ('--' + params['boundary']).encode('utf-8')
    for raw_part in body.split(boundary):
        part = raw_part.strip(b'\r\n')
        if not part or part == b'--':
            continue
        if part.endswith(b'--'):
            part = part[:-2].rstrip(b'\r\n')
        if b'\r\n\r\n' not in part:
            continue
        header_bytes, content = part.split(b'\r\n\r\n', 1)
        headers = {}
        for line in header_bytes.decode('utf-8', errors='replace').split('\r\n'):
            if ':' in line:
                k, v = line.split(':', 1)
                headers[k.strip().lower()] = v.strip()
        disp = parse_content_disposition(headers.get('content-disposition', ''))
        if disp.get('name') == 'file' and disp.get('filename'):
            return disp['filename'], content.rstrip(b'\r\n')
    raise ValueError('Không thấy trường file trong multipart request')

def sanitize_filename(name: str) -> str:
    name = re.sub(r"[^\w.()\-\sÀ-ỹ]", "_", name, flags=re.UNICODE).strip()
    return name[:160] or "upload.bin"


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".md", ".csv", ".json", ".html"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    return f"[Chưa hỗ trợ định dạng {suffix}. Hãy upload PDF, DOCX, TXT, MD, CSV hoặc JSON.]"


def extract_pdf(path: Path) -> str:
    try:
        import pypdf  # type: ignore
        reader = pypdf.PdfReader(str(path))
        chunks = []
        for idx, page in enumerate(reader.pages):
            try:
                chunks.append(f"\n\n--- PAGE {idx + 1} ---\n" + (page.extract_text() or ""))
            except Exception as exc:
                chunks.append(f"\n\n--- PAGE {idx + 1} ERROR: {exc} ---")
        return "\n".join(chunks).strip()
    except Exception as exc:
        return f"[Không trích xuất được PDF. Cài pypdf bằng: pip install pypdf. Lỗi: {exc}]"


def extract_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
        doc = docx.Document(str(path))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception as exc:
        return f"[Không trích xuất được DOCX. Cài python-docx bằng: pip install python-docx. Lỗi: {exc}]"


def extract_urls(text: str) -> list[str]:
    urls = re.findall(r"https?://[^\s)\]>'\"]+", text)
    seen, out = set(), []
    for u in urls:
        u = u.rstrip(".,;:")
        if u not in seen:
            out.append(u)
            seen.add(u)
        if len(out) >= 50:
            break
    return out


def classify_text(text: str) -> list[str]:
    low = text.lower()
    matches = []
    for section, keys in SECTION_KEYWORDS.items():
        if any(k.lower() in low for k in keys):
            matches.append(section)
    return matches or ["unclassified"]


def split_sentences(text: str) -> list[str]:
    clean = re.sub(r"\s+", " ", text).strip()
    # Vietnamese punctuation + Latin punctuation.
    sentences = re.split(r"(?<=[.!?。])\s+", clean)
    return [s.strip() for s in sentences if len(s.strip()) >= 60]


def sentence_score(sentence: str) -> float:
    low = sentence.lower()
    score = min(len(sentence) / 220, 1.0)
    for k in SUMMARY_KEYWORDS:
        if k.lower() in low:
            score += 1
    # Prefer sentences with numbers because market reports need metrics.
    if re.search(r"\d", sentence):
        score += 0.6
    return score


def summarize(text: str, max_sentences: int = 4) -> str:
    sentences = split_sentences(text)[:120]
    if not sentences:
        return text[:700]
    ranked = sorted(sentences, key=sentence_score, reverse=True)[:max_sentences]
    return " ".join(ranked)


def infer_title(filename: str, text: str) -> str:
    for line in text.splitlines()[:30]:
        line = line.strip(" #\t\r")
        if 12 <= len(line) <= 120 and not line.lower().startswith("--- page"):
            return line
    return filename


def append_upload(record: dict) -> None:
    try:
        existing = json.loads(USER_UPLOADS.read_text(encoding="utf-8")) if USER_UPLOADS.exists() else []
        if not isinstance(existing, list):
            existing = []
    except Exception:
        existing = []
    existing.insert(0, record)
    USER_UPLOADS.write_text(json.dumps(existing[:500], ensure_ascii=False, indent=2), encoding="utf-8")


class Handler(SimpleHTTPRequestHandler):
    def translate_path(self, path: str) -> str:
        parsed = urlparse(path)
        rel = parsed.path.lstrip("/") or "index.html"
        return str((BASE / rel).resolve())

    def do_POST(self):
        if self.path != "/api/ingest":
            self.send_error(404, "Not found")
            return
        try:
            self.handle_ingest()
        except Exception as exc:
            self.send_response(500)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(exc)}, ensure_ascii=False).encode("utf-8"))

    def handle_ingest(self):
        filename, content = read_multipart_file(self)
        original = sanitize_filename(filename)
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        saved = UPLOAD_DIR / f"{stamp}-{original}"
        with saved.open("wb") as f:
            f.write(content)
        text = extract_text(saved)
        urls = extract_urls(text)
        sections = classify_text(text)
        record = {
            "id": f"upload-{int(time.time()*1000)}",
            "filename": original,
            "saved_file": str(saved.relative_to(BASE)),
            "title": infer_title(original, text),
            "created_at": datetime.now(timezone.utc).isoformat(),
            "detected_sections": sections,
            "summary": summarize(text),
            "extracted_sources": urls,
            "char_count": len(text),
            "text_preview": text[:5000],
            "mode": "server-ingest"
        }
        append_upload(record)
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.end_headers()
        self.wfile.write(json.dumps(record, ensure_ascii=False).encode("utf-8"))


def main():
    port = int(os.environ.get("PORT", "8000"))
    os.chdir(BASE)
    server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    print(f"Bản đồ thị trường AI tác nhân đang chạy tại http://127.0.0.1:{port}")
    print("Nhấn Ctrl+C để dừng.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopped.")


if __name__ == "__main__":
    main()
